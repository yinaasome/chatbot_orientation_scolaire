import os
from dotenv import load_dotenv
import re
import json
import datetime
import unicodedata
import difflib
from typing import Dict, List, Optional, Tuple

import pandas as pd
import streamlit as st
from pymongo import MongoClient
from bson import ObjectId
import gridfs
import bcrypt

from google import genai

try:
    import requests
    from bs4 import BeautifulSoup
    _WEB_DISPONIBLE = True
except ImportError:
    _WEB_DISPONIBLE = False

load_dotenv()
# =========================================================================
# ⚙️ CONFIGURATION — adapte ces valeurs si besoin (ou via variables d'env)
# =========================================================================

DOCS_FOLDER = os.environ.get(
    "EDUCAT_FOLDER",
    "./data"
)

GEMINI_MODEL = os.environ.get(
    "GEMINI_CHAT_MODEL",
    "gemini-3.6-flash"
)

MONGO_URI = os.environ.get("MONGO_URI")

GOOGLE_SEARCH_API_KEY = os.environ.get(
    "GOOGLE_SEARCH_API_KEY"
)

GOOGLE_CSE_ID = os.environ.get(
    "GOOGLE_CSE_ID"
)

# =========================
# VERIFICATIONS
# =========================

if not MONGO_URI:
    raise ValueError(
        "❌ MONGO_URI n'est pas définie"
    )

if not GOOGLE_SEARCH_API_KEY:
    raise ValueError(
        "❌ GOOGLE_SEARCH_API_KEY n'est pas définie"
    )

if not GOOGLE_CSE_ID:
    raise ValueError(
        "❌ GOOGLE_CSE_ID n'est pas défini"
    )

# =========================
# MONGODB
# =========================

try:
    client = MongoClient(
        MONGO_URI,
        serverSelectionTimeoutMS=5000
    )

    client.admin.command("ping")

    print("✅ Connexion à MongoDB Atlas réussie !")

except Exception as e:
    print("❌ Erreur MongoDB :")
    print(e)

# Extensions reconnues
EXT_TABULAIRE = {"xlsx", "xls", "ods", "csv"}
EXT_DOCUMENT = {"pdf", "docx"}

# Sites officiels autorisés pour la recherche web complémentaire
SITES_OFFICIELS = {
    "CampusFaso":  "https://www.campusfaso.bf",
    "MESRSI":      "https://www.mesrsi.gov.bf",
    "MENA":        "http://www.mena.gov.bf/",
    "CIO-SPB":     "http://www.ciospb.bf/",
    "FONER":       "http://foner-bf.com/",
    "CENOU":       "https://www.cenou.gov.bf/",
    "oreille de campus faso": "https://www.oreilleducampus.org/",
    "le grand frere": "https://legrandfrere.africa/",
}

# Mots-clés qui orientent vers tel ou tel site officiel quand une
# recherche web de repli (sans API) est nécessaire
MOTS_CLES_SITES = {
    "CampusFaso": ["orientation", "affectation", "post-bac", "postbac", "campusfaso",
                   "inscription en ligne", "plateforme", "choix de filiere", "voeux"],
    "MESRSI":     ["ministere", "mesrsi", "enseignement superieur", "reglementation",
                   "texte", "decret", "arrete", "habilitation", "accreditation"],
    "MENA":       ["mena", "secondaire", "baccalaureat", "bac", "examen", "lycee", "college"],
    "CIO-SPB":    ["cio", "conseiller d'orientation", "conseil d'orientation", "psychologue scolaire"],
    "FONER":      ["bourse", "pret", "foner", "aide financiere", "financement des etudes"],
    "CENOU":      ["cite universitaire", "cenou", "logement etudiant", "restaurant universitaire",
                   "bourse d'etudes", "oeuvres universitaires"],
}

st.set_page_config(page_title="🎓 Orientation Universitaire", layout="wide")


# =========================================================================
# 🔡 Normalisation & tokenisation (base du moteur de recherche croisée)
# =========================================================================
_MOTS_VIDES = {
    "le", "la", "les", "de", "des", "du", "un", "une", "et", "ou", "est", "dans",
    "pour", "que", "qui", "avec", "sur", "au", "aux", "en", "a", "l", "d",
    "ce", "cette", "ces", "cet", "quel", "quelle", "quels", "quelles", "comment",
    "quoi", "donne", "moi", "je", "tu", "il", "elle", "ils", "elles", "peux",
    "peut", "voudrais", "sont", "son", "sa", "ses", "mon", "ma", "mes", "ton",
    "ta", "tes", "notre", "nos", "votre", "vos", "leur", "leurs", "y", "se",
    "qu", "as-tu", "avez", "vous", "nous", "on", "si", "car", "mais", "donc",
    "or", "ni", "par", "pas", "plus", "aussi", "etc",
}


def normaliser_texte(texte) -> str:
    """Minuscule + suppression des accents, pour comparer 'Université'
    et 'universite' ou 'Ki-Zerbo' et 'ki zerbo' sans faux négatifs."""
    s = str(texte).lower().strip()
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    return s


def tokeniser(texte) -> List[str]:
    s = normaliser_texte(texte)
    mots = re.findall(r"[a-z0-9]+", s)
    return [m for m in mots if len(m) > 1 and m not in _MOTS_VIDES]


def meilleure_similarite(a: str, b: str) -> float:
    """Similarité floue (0 à 1) tolérante aux fautes de frappe / variantes."""
    return difflib.SequenceMatcher(None, normaliser_texte(a), normaliser_texte(b)).ratio()


# =========================================================================
# 🔧 MongoDB
# =========================================================================
@st.cache_resource
def init_mongodb():
    try:
        client = MongoClient(MONGO_URI, serverSelectionTimeoutMS=5000)
        client.admin.command("ping")
        db = client["chatbot_orientation"]
        return db
    except Exception as e:
        st.error(f"Erreur de connexion à MongoDB : {e}")
        return None


# =========================================================================
# 🔐 Authentification
# =========================================================================
class AuthManager:
    def __init__(self, db):
        self.db = db
        self.users = db.users

    def hash_password(self, password: str) -> str:
        return bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")

    def verify_password(self, password: str, hashed: str) -> bool:
        return bcrypt.checkpw(password.encode("utf-8"), hashed.encode("utf-8"))

    def create_user(self, username: str, password: str, email: str = "", role: str = "Étudiant") -> bool:
        if self.users.find_one({"username": username}):
            return False
        self.users.insert_one({
            "username": username,
            "password": self.hash_password(password),
            "email": email,
            "role": role,
            "created_at": datetime.datetime.now(),
            "last_login": None,
            "is_active": True,
        })
        return True

    def authenticate_user(self, username: str, password: str) -> Optional[Dict]:
        user = self.users.find_one({"username": username, "is_active": True})
        if user and self.verify_password(password, user["password"]):
            self.users.update_one({"_id": user["_id"]}, {"$set": {"last_login": datetime.datetime.now()}})
            return {
                "user_id": str(user["_id"]),
                "username": user["username"],
                "email": user["email"],
                "role": user["role"],
            }
        return None


# =========================================================================
# 🧑‍🎓 Profil étudiant enrichi — aspirations + bulletins + relevé de bac
# =========================================================================
QUESTIONS_ASPIRATIONS = [
    {"cle": "domaines_interet", "label": "Domaines qui t'intéressent le plus", "type": "multiselect",
     "options": ["Sciences & Ingénierie", "Santé", "Droit & Sciences politiques", "Économie & Gestion",
                 "Lettres & Langues", "Sciences humaines & sociales", "Agriculture & Environnement",
                 "Arts & Communication", "Numérique & Informatique", "Enseignement"]},
    {"cle": "metiers_envisages", "label": "Quel(s) métier(s) envisages-tu (même de façon incertaine) ?", "type": "texte"},
    {"cle": "serie_bac", "label": "Série du bac", "type": "select",
     "options": ["A1", "A2", "A3", "A4", "C", "D", "E", "F1", "F2", "F3", "F4", "G1", "G2", "G3", "Autre"]},
    {"cle": "mention_bac", "label": "Mention obtenue (ou espérée) au bac", "type": "select",
     "options": ["Passable", "Assez Bien", "Bien", "Très Bien", "Excellent", "Je ne sais pas encore"]},
    {"cle": "annee_bac", "label": "Année d'obtention du bac", "type": "texte"},
    {"cle": "zone_obtention_bac", "label": "Région / zone où tu as obtenu (ou passeras) le bac", "type": "select",
     "options": ["Boucle du Mouhoun", "Cascades", "Centre", "Centre-Est", "Centre-Nord", "Centre-Ouest",
                 "Centre-Sud", "Est", "Hauts-Bassins", "Nord", "Plateau-Central", "Sahel", "Sud-Ouest"]},
    {"cle": "ville_residence", "label": "Ville où tu résides actuellement", "type": "texte"},
    {"cle": "contraintes", "label": "Contraintes ou préférences à prendre en compte", "type": "multiselect",
     "options": ["Budget familial limité", "Souhaite rester proche de chez soi", "Prêt(e) à changer de ville",
                 "Besoin d'une cité universitaire / logement", "Souhaite une bourse ou une aide financière",
                 "Préfère le public", "Préfère le privé", "Aucune contrainte particulière"]},
]

DOCUMENTS_REQUIS = [
    {"cle": "bulletin_t1", "label": "Bulletin de terminale — 1er trimestre"},
    {"cle": "bulletin_t2", "label": "Bulletin de terminale — 2e trimestre"},
    {"cle": "bulletin_t3", "label": "Bulletin de terminale — 3e trimestre"},
    {"cle": "releve_bac", "label": "Relevé de notes du baccalauréat"},
]


class StudentProfileManager:
    """Gère les aspirations, les bulletins/relevé (stockés dans GridFS) et
    l'analyse académique du profil de chaque étudiant."""

    def __init__(self, db):
        self.db = db
        self.profiles = db.student_profiles
        self.fs = gridfs.GridFS(db, collection="bulletins_fichiers")

    def get_profile(self, user_id: str) -> Dict:
        return self.profiles.find_one({"user_id": user_id}) or {}

    def save_aspirations(self, user_id: str, reponses: Dict):
        self.profiles.update_one(
            {"user_id": user_id},
            {"$set": {"aspirations": reponses, "aspirations_completees": True,
                      "updated_at": datetime.datetime.now()}},
            upsert=True,
        )

    def sauvegarder_document(self, user_id: str, cle_document: str, nom_fichier: str,
                              contenu_bytes: bytes, texte_extrait: str):
        """Stocke le PDF original dans GridFS et le texte extrait + métadonnées
        dans student_profiles.documents.<cle_document>."""
        # On retire l'éventuel ancien fichier du même type pour ne pas accumuler
        profil = self.get_profile(user_id)
        ancien = (profil.get("documents") or {}).get(cle_document, {})
        ancien_id = ancien.get("gridfs_id")
        if ancien_id:
            try:
                self.fs.delete(ObjectId(ancien_id))
            except Exception:
                pass

        gridfs_id = self.fs.put(contenu_bytes, filename=nom_fichier, user_id=user_id, type_document=cle_document)
        self.profiles.update_one(
            {"user_id": user_id},
            {"$set": {
                f"documents.{cle_document}": {
                    "nom_fichier": nom_fichier,
                    "gridfs_id": str(gridfs_id),
                    "texte_extrait": texte_extrait,
                    "charge_le": datetime.datetime.now(),
                },
                "updated_at": datetime.datetime.now(),
            }},
            upsert=True,
        )

    def documents_manquants(self, user_id: str) -> List[str]:
        profil = self.get_profile(user_id)
        documents = profil.get("documents") or {}
        return [d["cle"] for d in DOCUMENTS_REQUIS if d["cle"] not in documents]

    def profil_complet(self, user_id: str) -> bool:
        profil = self.get_profile(user_id)
        return bool(profil.get("aspirations_completees")) and not self.documents_manquants(user_id)

    def save_analyse_academique(self, user_id: str, analyse: Dict):
        self.profiles.update_one(
            {"user_id": user_id},
            {"$set": {"analyse_academique": analyse, "updated_at": datetime.datetime.now()}},
            upsert=True,
        )

    def analyser_bulletins_avec_llm(self, gemini_client, user_id: str) -> Dict:
        """Envoie les textes extraits des bulletins + relevé de bac au LLM pour
        en tirer un profil académique structuré (matières fortes/faibles, etc.)"""
        profil = self.get_profile(user_id)
        documents = profil.get("documents") or {}
        textes = []
        for doc in DOCUMENTS_REQUIS:
            info = documents.get(doc["cle"])
            if info and info.get("texte_extrait"):
                textes.append(f"### {doc['label']} ({info['nom_fichier']})\n{info['texte_extrait'][:6000]}")

        if not textes:
            return {"erreur": "Aucun bulletin exploitable n'a été trouvé."}

        prompt_analyse = f"""
Tu es un conseiller d'orientation scolaire. Voici le texte extrait de bulletins de
classe de terminale et/ou d'un relevé de notes du baccalauréat d'un(e) élève
burkinabè. Le texte peut être imparfait (issu d'une extraction automatique de PDF).

TEXTE DES BULLETINS / RELEVÉ :
{chr(10).join(textes)}

Analyse ces documents et réponds UNIQUEMENT avec un objet JSON valide (rien
d'autre, pas de balises ```), avec exactement cette structure :
{{
  "moyenne_generale_estimee": "ex: 13/20 (estimation) ou 'non déterminable'",
  "matieres_fortes": ["liste de matières où l'élève réussit le mieux"],
  "matieres_faibles": ["liste de matières où l'élève est en difficulté"],
  "profil_dominant": "un mot parmi: scientifique, littéraire, économique, technique, mixte",
  "regularite": "brève phrase sur la progression/régularité si visible sur plusieurs trimestres",
  "recommandation_synthetique": "2-3 phrases de synthèse utile pour orienter l'élève"
}}
Si une information est absente ou illisible, indique-le honnêtement (ex: "non déterminable") plutôt que d'inventer.
"""
        try:
            reponse = gemini_client.models.generate_content(model=GEMINI_MODEL, contents=prompt_analyse)
            texte = (reponse.text or "").strip()
            texte = re.sub(r"^```(json)?|```$", "", texte, flags=re.MULTILINE).strip()
            analyse = json.loads(texte)
        except Exception as e:
            analyse = {"erreur": f"Analyse automatique impossible ({e}).",
                       "recommandation_synthetique": "Analyse manuelle recommandée."}

        self.save_analyse_academique(user_id, analyse)
        return analyse

    def construire_contexte_profil(self, user_id: str) -> str:
        """Résumé texte du profil enrichi, injecté dans le prompt principal."""
        profil = self.get_profile(user_id)
        if not profil:
            return "(profil non renseigné — recommandations génériques uniquement)"

        blocs = []
        aspirations = profil.get("aspirations")
        if aspirations:
            lignes = []
            for q in QUESTIONS_ASPIRATIONS:
                val = aspirations.get(q["cle"])
                if val:
                    val_txt = ", ".join(val) if isinstance(val, list) else val
                    lignes.append(f"- {q['label']} : {val_txt}")
            blocs.append("ASPIRATIONS ET SITUATION DE L'ÉLÈVE :\n" + "\n".join(lignes))

        analyse = profil.get("analyse_academique")
        if analyse and not analyse.get("erreur"):
            blocs.append(
                "PROFIL ACADÉMIQUE (issu de l'analyse des bulletins/relevé de bac) :\n"
                f"- Moyenne générale estimée : {analyse.get('moyenne_generale_estimee', 'n/d')}\n"
                f"- Matières fortes : {', '.join(analyse.get('matieres_fortes', [])) or 'n/d'}\n"
                f"- Matières faibles : {', '.join(analyse.get('matieres_faibles', [])) or 'n/d'}\n"
                f"- Profil dominant : {analyse.get('profil_dominant', 'n/d')}\n"
                f"- Régularité : {analyse.get('regularite', 'n/d')}\n"
                f"- Synthèse : {analyse.get('recommandation_synthetique', 'n/d')}"
            )
        elif analyse and analyse.get("erreur"):
            blocs.append(f"PROFIL ACADÉMIQUE : {analyse['erreur']}")

        return "\n\n".join(blocs) if blocs else "(profil non renseigné — recommandations génériques uniquement)"


# =========================================================================
# 🗃️ Mémoire des conversations / contexte utilisateur
# =========================================================================
class ChatbotMemory:
    def __init__(self, db):
        self.db = db
        self.conversations = db.conversations
        self.user_context = db.user_context
        self.chat_sessions = db.chat_sessions

    def create_new_chat_session(self, user_id: str) -> str:
        result = self.chat_sessions.insert_one({
            "user_id": user_id,
            "created_at": datetime.datetime.now(),
            "title": f"Chat du {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}",
        })
        return str(result.inserted_id)

    def get_user_chat_sessions(self, user_id: str, limit: int = 10) -> List[Dict]:
        return list(self.chat_sessions.find({"user_id": user_id}).sort("created_at", -1).limit(limit))

    def get_session_conversations(self, session_id: str) -> List[Dict]:
        return list(self.conversations.find({"session_id": session_id}).sort("timestamp", 1))

    def save_conversation(self, user_id: str, session_id: str, question: str, response: str, metadata: Dict = None):
        return self.conversations.insert_one({
            "user_id": user_id,
            "session_id": session_id,
            "question": question,
            "response": response,
            "timestamp": datetime.datetime.now(),
            "metadata": metadata or {},
        })

    def get_conversation_history(self, user_id: str, limit: int = 5):
        return list(self.conversations.find({"user_id": user_id}).sort("timestamp", -1).limit(limit))

    def get_user_context(self, user_id: str):
        return self.user_context.find_one({"user_id": user_id}) or {}

    def update_user_context(self, user_id: str, context_data: Dict):
        self.user_context.update_one(
            {"user_id": user_id}, {"$set": {**context_data, "last_updated": datetime.datetime.now()}}, upsert=True
        )

    def add_recent_entity(self, user_id: str, entity_type: str, entity_name: str):
        context = self.get_user_context(user_id)
        key = f"recent_{entity_type}"
        items = context.get(key, [])
        items = [i for i in items if i["name"] != entity_name]
        items.insert(0, {"name": entity_name, "timestamp": datetime.datetime.now()})
        context[key] = items[:20]
        self.update_user_context(user_id, context)

    def update_session_title(self, session_id: str, title: str):
        self.chat_sessions.update_one({"_id": ObjectId(session_id)}, {"$set": {"title": title}})


# =========================================================================
# 📚 Ingestion générique de TOUTES les sources du dossier Educat
# =========================================================================
def _signature_dossier(dossier: str) -> str:
    """Empreinte du dossier (noms + tailles + dates de modif.) pour invalider
    automatiquement le cache dès qu'un fichier est ajouté / modifié / supprimé —
    c'est ce qui permet d'ajouter un nouveau PDF/Word/Excel sans toucher au code."""
    empreinte = []
    if not os.path.isdir(dossier):
        return "dossier_absent"
    for racine, _, fichiers in os.walk(dossier):
        for f in sorted(fichiers):
            chemin = os.path.join(racine, f)
            try:
                stat = os.stat(chemin)
                empreinte.append(f"{chemin}:{stat.st_size}:{stat.st_mtime}")
            except OSError:
                pass
    return "|".join(empreinte)


@st.cache_data(show_spinner="📊 Chargement des données structurées (Excel / ODS / CSV)...")
def charger_toutes_les_tables(dossier: str, _signature: str) -> Dict[str, pd.DataFrame]:
    """Charge CHAQUE feuille de CHAQUE fichier tabulaire (.ods/.xlsx/.xls/.csv)
    trouvé dans le dossier — et pas uniquement un fichier fixe codé en dur.
    Clé du dictionnaire : "nom_du_fichier :: nom_de_la_feuille"."""
    tables: Dict[str, pd.DataFrame] = {}
    if not os.path.isdir(dossier):
        return tables

    for racine, _, fichiers in os.walk(dossier):
        for nom_fichier in fichiers:
            chemin = os.path.join(racine, nom_fichier)
            ext = nom_fichier.lower().rsplit(".", 1)[-1] if "." in nom_fichier else ""
            if ext not in EXT_TABULAIRE:
                continue
            try:
                if ext == "csv":
                    df = pd.read_csv(chemin)
                    if not df.empty:
                        tables[nom_fichier] = df
                else:
                    engine = "odf" if ext == "ods" else None
                    xls = pd.ExcelFile(chemin, engine=engine)
                    for feuille in xls.sheet_names:
                        df = pd.read_excel(xls, sheet_name=feuille)
                        if not df.empty:
                            tables[f"{nom_fichier} :: {feuille}"] = df
            except Exception as e:
                tables[f"{nom_fichier} (ERREUR DE LECTURE)"] = pd.DataFrame(
                    {"erreur": [f"Impossible de lire ce fichier : {e}"]}
                )
    return tables


OCR_PAGES_MAX = int(os.environ.get("OCR_PAGES_MAX", "20"))  # borne de sécurité (temps/mémoire)
OCR_DPI = int(os.environ.get("OCR_DPI", "150"))
OCR_TAILLE_MAX_IMAGE = (1600, 1600)  # redimensionnement avant OCR (mémoire/vitesse)


def _extraire_texte_pdf(chemin: str) -> List[Tuple[int, str]]:
    """Retourne une liste (numéro_page, texte) pour permettre de citer la page
    source. Tente un repli OCR (pytesseract) si le PDF est scanné / sans texte
    extractible et que les librairies optionnelles sont disponibles.

    L'OCR est fait PAGE PAR PAGE (jamais tout le PDF d'un coup) et les images
    sont redimensionnées avant reconnaissance, pour éviter de saturer la
    mémoire sur de gros PDF scannés haute résolution comme "Guide_orientation.pdf".
    Une page qui échoue (langue Tesseract manquante, page corrompue...) est
    simplement signalée, sans faire échouer le reste du document."""
    pages_texte: List[Tuple[int, str]] = []
    try:
        from pypdf import PdfReader
        reader = PdfReader(chemin)
        for i, page in enumerate(reader.pages, start=1):
            texte = (page.extract_text() or "").strip()
            pages_texte.append((i, texte))
    except Exception as e:
        return [(0, f"[Erreur de lecture PDF : {e}]")]

    texte_total = sum(len(t) for _, t in pages_texte)
    nb_pages = max(len(pages_texte), 1)
    if texte_total >= 20 * nb_pages:
        return pages_texte  # texte déjà exploitable, pas besoin d'OCR

    # Très peu / pas de texte extrait -> probablement un PDF scanné (image).
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError:
        pages_texte.append((0, "[Document probablement scanné : aucun texte n'a pu être "
                                "extrait directement. Installez 'pytesseract' + 'pdf2image' "
                                "(+ le binaire Tesseract-OCR et son paquet de langue "
                                "'tesseract-ocr-fra') pour permettre l'OCR de ce fichier.]"))
        return pages_texte

    pages_a_traiter = min(nb_pages, OCR_PAGES_MAX)
    pages_ocr: List[Tuple[int, str]] = []
    for i in range(1, pages_a_traiter + 1):
        try:
            image = convert_from_path(chemin, dpi=OCR_DPI, first_page=i, last_page=i)[0]
            image.thumbnail(OCR_TAILLE_MAX_IMAGE)
            texte_page = pytesseract.image_to_string(image, lang="fra")
            pages_ocr.append((i, texte_page.strip()))
            del image
        except Exception as e:
            pages_ocr.append((i, f"[OCR impossible sur cette page : {e}]"))

    if nb_pages > OCR_PAGES_MAX:
        pages_ocr.append((0, f"[Seules les {OCR_PAGES_MAX} premières pages sur {nb_pages} "
                              f"ont été passées à l'OCR (limite OCR_PAGES_MAX). Augmentez "
                              f"cette variable d'environnement si besoin.]"))

    if sum(len(t) for _, t in pages_ocr) > texte_total:
        return pages_ocr
    return pages_texte


def _extraire_texte_docx(chemin: str) -> str:
    try:
        import docx
        d = docx.Document(chemin)
        morceaux = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                morceaux.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(m for m in morceaux if m.strip())
    except Exception as e:
        return f"[Erreur de lecture DOCX : {e}]"


@st.cache_data(show_spinner="📄 Analyse des documents PDF / Word du dossier...")
def construire_corpus_documents(dossier: str, _signature: str, taille_min: int = 40) -> List[Dict]:
    """Parcourt le dossier, extrait le texte de CHAQUE .pdf / .docx trouvé
    (peu importe son nom) et le découpe en paragraphes exploitables pour la
    recherche. Chaque chunk garde une trace de son fichier (et sa page pour
    les PDF), ce qui permet au LLM de citer sa source."""
    chunks: List[Dict] = []
    if not os.path.isdir(dossier):
        return chunks

    for racine, _, fichiers in os.walk(dossier):
        for nom_fichier in fichiers:
            chemin = os.path.join(racine, nom_fichier)
            ext = nom_fichier.lower().rsplit(".", 1)[-1] if "." in nom_fichier else ""

            if ext == "pdf":
                for numero_page, texte_page in _extraire_texte_pdf(chemin):
                    for paragraphe in re.split(r"\n\s*\n", texte_page):
                        paragraphe = re.sub(r"\s+", " ", paragraphe).strip()
                        if len(paragraphe) >= taille_min:
                            chunks.append({"fichier": nom_fichier, "page": numero_page, "texte": paragraphe})

            elif ext == "docx":
                contenu = _extraire_texte_docx(chemin)
                for paragraphe in re.split(r"\n\s*\n", contenu):
                    paragraphe = re.sub(r"\s+", " ", paragraphe).strip()
                    if len(paragraphe) >= taille_min:
                        chunks.append({"fichier": nom_fichier, "page": None, "texte": paragraphe})
    return chunks


# =========================================================================
# 🔍 Moteur de recherche croisée — cœur de la correction demandée
# =========================================================================
def _ligne_en_texte(row: pd.Series) -> str:
    return " | ".join(f"{col} : {val}" for col, val in row.items() if pd.notna(val) and str(val).strip())


SCORE_MIN_TABLE = 1.0  # en dessous, on considère la correspondance comme du bruit (ex: feuilles README)


def rechercher_dans_toutes_les_tables(question: str, tables: Dict[str, pd.DataFrame],
                                       max_par_source: int = 6, max_total: int = 24) -> List[Tuple[float, str, str]]:
    """Cherche la question dans TOUTES les feuilles de TOUS les fichiers
    tabulaires, sans s'arrêter à la première source qui matche. Retourne les
    meilleurs résultats en garantissant une diversité de sources (au plus
    `max_par_source` lignes par fichier/feuille) plutôt qu'un unique tableau
    qui monopoliserait toute la réponse."""
    tokens_question = set(tokeniser(question))
    if not tokens_question or not tables:
        return []

    resultats_par_source: Dict[str, List[Tuple[float, str]]] = {}

    for source, df in tables.items():
        candidats = []
        for _, row in df.iterrows():
            texte_ligne = _ligne_en_texte(row)
            if not texte_ligne:
                continue
            tokens_ligne = set(tokeniser(texte_ligne))
            intersection = tokens_question & tokens_ligne
            if not intersection:
                continue
            # Score = recouvrement des mots-clés, normalisé par la longueur de la question
            score = len(intersection) / (len(tokens_question) ** 0.5)
            # Bonus si un des mots de la question colle de très près (fautes de
            # frappe / variantes) à une valeur importante de la ligne, ex: le
            # nom de l'université ou de la filière
            for valeur in row.values:
                valeur_str = str(valeur)
                if len(valeur_str) > 3 and meilleure_similarite(valeur_str, question) > 0.5:
                    score += 1.5
                    break
            if score >= SCORE_MIN_TABLE:
                candidats.append((score, texte_ligne))
        candidats.sort(key=lambda x: x[0], reverse=True)
        if candidats:
            resultats_par_source[source] = candidats[:max_par_source]

    # Fusion en respectant la diversité de sources : on alterne entre sources
    # plutôt que de prendre les 24 meilleures lignes d'une seule feuille
    fusion: List[Tuple[float, str, str]] = []
    index_par_source = {s: 0 for s in resultats_par_source}
    sources_triees = sorted(resultats_par_source.keys(),
                             key=lambda s: resultats_par_source[s][0][0], reverse=True)
    while len(fusion) < max_total and any(
        index_par_source[s] < len(resultats_par_source[s]) for s in sources_triees
    ):
        for source in sources_triees:
            i = index_par_source[source]
            if i < len(resultats_par_source[source]):
                score, texte = resultats_par_source[source][i]
                fusion.append((score, source, texte))
                index_par_source[source] += 1
            if len(fusion) >= max_total:
                break

    fusion.sort(key=lambda x: x[0], reverse=True)
    return fusion[:max_total]


def rechercher_dans_les_documents(question: str, chunks: List[Dict], top_k: int = 8) -> List[Dict]:
    """Recherche par recouvrement de tokens normalisés sur tous les chunks de
    tous les PDF/DOCX du dossier (quel que soit leur nombre ou leur nom)."""
    tokens_question = set(tokeniser(question))
    if not tokens_question or not chunks:
        return []

    resultats = []
    for chunk in chunks:
        tokens_chunk = set(tokeniser(chunk["texte"]))
        intersection = tokens_question & tokens_chunk
        if not intersection:
            continue
        score = len(intersection) / (len(tokens_question) ** 0.5)
        resultats.append((score, chunk))
    resultats.sort(key=lambda x: x[0], reverse=True)
    return [c for _, c in resultats[:top_k]]


def formater_resultats_tables(resultats: List[Tuple[float, str, str]]) -> str:
    if not resultats:
        return "(aucune correspondance trouvée dans les tableaux de données)"
    blocs = {}
    for score, source, texte in resultats:
        blocs.setdefault(source, []).append(texte)
    morceaux = []
    for source, lignes in blocs.items():
        morceaux.append(f"### Source : {source}\n" + "\n".join(f"- {l}" for l in lignes))
    return "\n\n".join(morceaux)


def formater_resultats_documents(resultats: List[Dict]) -> str:
    if not resultats:
        return "(aucun extrait pertinent trouvé dans les documents PDF/Word)"
    morceaux = []
    for chunk in resultats:
        reference = chunk["fichier"] if chunk.get("page") in (None, 0) else f"{chunk['fichier']}, page {chunk['page']}"
        morceaux.append(f"[{reference}]\n{chunk['texte']}")
    return "\n\n".join(morceaux)


def recuperer_toutes_les_valeurs_connues(tables: Dict[str, pd.DataFrame]) -> List[str]:
    """Rassemble les noms distincts (universités, filières, cités...) pour le
    suivi du contexte utilisateur (entités récemment consultées), toutes
    sources confondues."""
    valeurs = set()
    for df in tables.values():
        for col in df.select_dtypes(include="object").columns:
            for v in df[col].dropna().astype(str).unique():
                v = v.strip()
                if 3 < len(v) < 80:
                    valeurs.add(v)
    return list(valeurs)


# =========================================================================
# 🌐 Recherche complémentaire sur les sites officiels
# =========================================================================
def _site_pertinent_pour(question: str) -> List[str]:
    q_norm = normaliser_texte(question)
    sites_trouves = []
    for site, mots_cles in MOTS_CLES_SITES.items():
        if any(normaliser_texte(m) in q_norm for m in mots_cles):
            sites_trouves.append(site)
    return sites_trouves


def question_merite_recherche_web(question: str, meilleur_score_local: float) -> bool:
    """Décide si une vérification sur un site officiel est utile : soit parce
    que la question touche un thème couvert par un site officiel précis, soit
    parce que les données locales ne donnent pas une réponse solide."""
    if _site_pertinent_pour(question):
        return True
    if meilleur_score_local < 0.6:
        return True
    return False


@st.cache_data(ttl=1800, show_spinner=False)
def rechercher_sites_officiels(question: str) -> str:
    """Recherche complémentaire restreinte aux sites officiels listés.
    - Si GOOGLE_SEARCH_API_KEY / GOOGLE_CSE_ID sont configurées (Google
      Programmable Search Engine restreint à ces domaines), on fait une vraie
      recherche plein texte.
    - Sinon, on se contente d'aller chercher le texte des pages d'accueil des
      sites jugés pertinents pour la question (mode de repli, moins précis)."""
    if not _WEB_DISPONIBLE:
        return "(recherche web indisponible : installez 'requests' et 'beautifulsoup4')"

    sites_cibles = _site_pertinent_pour(question) or list(SITES_OFFICIELS.keys())
    resultats_textuels = []

    if GOOGLE_SEARCH_API_KEY and GOOGLE_CSE_ID:
        for site in sites_cibles:
            url_domaine = SITES_OFFICIELS[site].replace("https://", "").replace("http://", "").rstrip("/")
            try:
                reponse = requests.get(
                    "https://www.googleapis.com/customsearch/v1",
                    params={
                        "key": GOOGLE_SEARCH_API_KEY,
                        "cx": GOOGLE_CSE_ID,
                        "q": question,
                        "siteSearch": url_domaine,
                        "num": 3,
                    },
                    timeout=8,
                )
                donnees = reponse.json()
                for item in donnees.get("items", []):
                    resultats_textuels.append(
                        f"[{site} — {item.get('link')}]\n{item.get('title', '')} — {item.get('snippet', '')}"
                    )
            except Exception as e:
                resultats_textuels.append(f"[{site}] Recherche indisponible ({e})")
    else:
        # Mode de repli sans API : on récupère le texte visible de la page
        # d'accueil de chaque site jugé pertinent.
        for site in sites_cibles:
            url = SITES_OFFICIELS[site]
            try:
                reponse = requests.get(url, timeout=8, headers={"User-Agent": "Mozilla/5.0"})
                soupe = BeautifulSoup(reponse.text, "html.parser")
                for balise in soupe(["script", "style", "nav", "footer"]):
                    balise.decompose()
                texte = re.sub(r"\s+", " ", soupe.get_text(" ")).strip()
                resultats_textuels.append(f"[{site} — {url}]\n{texte[:1200]}")
            except Exception as e:
                resultats_textuels.append(f"[{site} — {url}] Page inaccessible ({e})")

    if not resultats_textuels:
        return "(aucun résultat officiel pertinent trouvé)"

    avertissement = (
        "" if (GOOGLE_SEARCH_API_KEY and GOOGLE_CSE_ID) else
        "\n(Note : recherche en mode de repli — page d'accueil uniquement, pas une "
        "recherche plein texte du site. Configurez GOOGLE_SEARCH_API_KEY et GOOGLE_CSE_ID "
        "pour une recherche officielle précise.)\n"
    )
    return avertissement + "\n\n".join(resultats_textuels)


# =========================================================================
# 🤖 Gemini
# =========================================================================
@st.cache_resource
def init_gemini_client():
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        st.error("⚠️ Variable d'environnement GEMINI_API_KEY manquante. Définis-la puis relance l'application.")
        return None
    try:
        return genai.Client(api_key=api_key)
    except Exception as e:
        st.error(f"Erreur d'initialisation de Gemini : {e}")
        return None


def generer_reponse(client, prompt: str) -> str:
    try:
        reponse = client.models.generate_content(model=GEMINI_MODEL, contents=prompt)
        return reponse.text or "Je n'ai pas pu générer de réponse."
    except Exception as e:
        return f"Erreur lors de l'appel à Gemini : {e}"


PROMPT_TEMPLATE = """
Tu es un conseiller d'orientation universitaire expérimenté au Burkina Faso. Tu
réponds à des questions sur les universités, UFR/instituts, filières,
conditions d'admission et cités universitaires (logement étudiant).

RÈGLES DE FOND :
- Base-toi en priorité sur les DONNÉES STRUCTURÉES et les EXTRAITS DE
  DOCUMENTS fournis ci-dessous. N'invente aucune information factuelle
  (chiffres, adresses, conditions précises) qui n'y figurerait pas.
- Ces données proviennent de PLUSIEURS sources différentes (fichiers Excel/CSV/
  PDF distincts) : CROISE-les activement. Si deux sources se complètent
  (ex : une donne la filière, une autre donne les débouchés ou la localisation),
  combine-les dans une réponse unique et cohérente plutôt que de t'arrêter à la
  première information trouvée. Si des sources se contredisent, signale-le.
- Si les données structurées et documentaires ne suffisent pas et qu'une
  RECHERCHE SUR LES SITES OFFICIELS est fournie plus bas, appuie-toi dessus en
  citant clairement le site et l'URL. Précise que ces informations méritent
  d'être reconfirmées sur le site car elles peuvent évoluer (frais, dates
  limites, procédures en ligne).
- Si une information manque vraiment partout, dis-le clairement et propose à
  l'élève de vérifier auprès de l'établissement ou du site officiel concerné,
  plutôt que d'inventer.

RÈGLES DE PÉDAGOGIE (très important) :
- Quand tu parles d'une UNIVERSITÉ : ne te limite pas à la lister. Explique
  aussi, à partir des données disponibles, les conditions réelles
  d'apprentissage — mode d'accès (orientation automatique post-bac ou
  concours/test d'entrée), type d'établissement (public/privé), UFR de
  rattachement, ville et donc conditions de vie/logement, langue
  d'enseignement si pertinent. Objectif : que l'élève comprenne concrètement
  à quoi ressemblera sa vie d'étudiant, pas seulement le nom de l'établissement.
- Quand tu parles d'une FILIÈRE : explique systématiquement les débouchés
  concrets (métiers réels accessibles), la tendance actuelle du marché de
  l'emploi pour cette filière au Burkina Faso (secteur qui recrute, secteur
  saturé, création d'entreprise possible, etc.), et donne 1 à 3 exemples
  concrets de métiers ou types d'employeurs. Le but est d'amener l'élève à
  comprendre ET à voir la réalité du terrain, pas juste un intitulé de diplôme.
- Adapte ton discours au PROFIL DE L'ÉLÈVE fourni plus bas (aspirations,
  contraintes, série du bac, points forts/faibles académiques) : mets en avant
  les options qui lui correspondent réellement et explique pourquoi.
- Si plusieurs options existent, présente-les de façon comparative
  (tableau ou liste à puces).
- Réponds de façon claire, structurée, professionnelle et bienveillante —
  jamais condescendante. Tu t'adresses à un(e) jeune qui doit prendre une
  décision importante pour son avenir.

🧑‍🎓 PROFIL ENRICHI DE L'ÉLÈVE :
{profil_eleve}

📋 CONTEXTE (universités / filières / cités déjà consultées) :
{contexte_utilisateur}

📚 HISTORIQUE RÉCENT DE LA CONVERSATION :
{historique}

📊 DONNÉES STRUCTURÉES (recherche croisée sur tous les fichiers Excel/CSV/ODS du dossier) :
{donnees_structurees}

📄 EXTRAITS DE DOCUMENTS (recherche croisée sur tous les PDF/Word du dossier) :
{documents_contextuels}

🌐 RECHERCHE SUR LES SITES OFFICIELS (le cas échéant) :
{recherche_web}

Question de l'utilisateur :
{question}

➡️ Donne une réponse professionnelle, bien structurée, et concrète.
"""


def construire_prompt(question, profil_eleve, donnees_structurees, documents_contextuels,
                       recherche_web, historique, contexte_utilisateur):
    return PROMPT_TEMPLATE.format(
        question=question,
        profil_eleve=profil_eleve or "(aucun profil renseigné)",
        donnees_structurees=donnees_structurees or "(aucune donnée structurée pertinente trouvée)",
        documents_contextuels=documents_contextuels or "(aucun document complémentaire pertinent trouvé)",
        recherche_web=recherche_web or "(non nécessaire pour cette question)",
        historique=historique or "(pas d'historique)",
        contexte_utilisateur=contexte_utilisateur or "(aucun)",
    )


# =========================================================================
# 🔁 Traitement d'une question
# =========================================================================
def traiter_question(question, tables, chunks_docs, valeurs_connues, gemini_client, memory,
                      profile_manager, user_id, session_id):
    # 1) Recherche croisée dans TOUTES les données structurées
    resultats_tables = rechercher_dans_toutes_les_tables(question, tables)
    donnees_structurees = formater_resultats_tables(resultats_tables)
    meilleur_score = resultats_tables[0][0] if resultats_tables else 0.0

    # 2) Recherche croisée dans TOUS les documents (pdf/docx)
    resultats_docs = rechercher_dans_les_documents(question, chunks_docs)
    documents_contextuels = formater_resultats_documents(resultats_docs)

    # 3) Recherche complémentaire sur les sites officiels si pertinent
    recherche_web = ""
    if question_merite_recherche_web(question, meilleur_score):
        recherche_web = rechercher_sites_officiels(question)

    # 4) Profil enrichi de l'élève
    profil_eleve = profile_manager.construire_contexte_profil(user_id)

    # 5) Historique et contexte utilisateur
    historique_brut = memory.get_conversation_history(user_id, 5)
    historique_text = ""
    for conv in reversed(historique_brut[:3]):
        historique_text += f"Q: {conv['question'][:100]}\nR: {conv['response'][:150]}\n\n"

    contexte = memory.get_user_context(user_id)
    contexte_text = ""
    for cle, etiquette in [("recent_universites", "Universités"), ("recent_filieres", "Filières"), ("recent_cites", "Cités")]:
        items = contexte.get(cle, [])
        if items:
            contexte_text += f"{etiquette} récemment consultées : {', '.join(i['name'] for i in items[:5])}\n"

    # 6) Appel à Gemini
    prompt = construire_prompt(question, profil_eleve, donnees_structurees, documents_contextuels,
                                recherche_web, historique_text, contexte_text)
    reponse = generer_reponse(gemini_client, prompt)

    # 7) Mémorisation des entités mentionnées (toutes sources confondues)
    question_lower = normaliser_texte(question)
    for valeur in valeurs_connues:
        if len(valeur) > 3 and normaliser_texte(valeur) in question_lower:
            valeur_lower = valeur.lower()
            if "universit" in valeur_lower or "institut" in valeur_lower or "ecole" in valeur_lower:
                memory.add_recent_entity(user_id, "universites", valeur)
            elif "cite" in valeur_lower:
                memory.add_recent_entity(user_id, "cites", valeur)

    memory.save_conversation(user_id, session_id, question, reponse, metadata={"sujet": "orientation"})
    return reponse


# =========================================================================
# 🔐 Page de connexion / inscription
# =========================================================================
def show_login_page(auth_manager):
    st.title("🎓 Orientation Universitaire — Connexion")
    tab1, tab2 = st.tabs(["🔑 Connexion", "👤 Inscription"])

    with tab1:
        with st.form("login_form"):
            username = st.text_input("Nom d'utilisateur")
            password = st.text_input("Mot de passe", type="password")
            if st.form_submit_button("Se connecter"):
                if username and password:
                    user = auth_manager.authenticate_user(username, password)
                    if user:
                        st.session_state.user = user
                        st.session_state.authenticated = True
                        st.rerun()
                    else:
                        st.error("Nom d'utilisateur ou mot de passe incorrect")
                else:
                    st.error("Veuillez saisir tous les champs")

    with tab2:
        with st.form("register_form"):
            new_username = st.text_input("Nom d'utilisateur", key="reg_username")
            new_email = st.text_input("Email (optionnel)", key="reg_email")
            new_password = st.text_input("Mot de passe", type="password", key="reg_password")
            confirm_password = st.text_input("Confirmer le mot de passe", type="password", key="reg_confirm")
            new_role = st.selectbox("Rôle", ["Étudiant", "Parent", "Conseiller", "Autre"], key="reg_role")
            if st.form_submit_button("Créer le compte"):
                if new_username and new_password and confirm_password:
                    if new_password != confirm_password:
                        st.error("Les mots de passe ne correspondent pas")
                    elif len(new_password) < 6:
                        st.error("Le mot de passe doit contenir au moins 6 caractères")
                    elif auth_manager.create_user(new_username, new_password, new_email, new_role):
                        st.success("Compte créé avec succès ! Vous pouvez maintenant vous connecter.")
                    else:
                        st.error("Ce nom d'utilisateur existe déjà")
                else:
                    st.error("Veuillez saisir tous les champs obligatoires")


# =========================================================================
# 🧑‍🎓 Assistant de constitution du profil enrichi (aspirations + bulletins)
# =========================================================================
def show_profile_onboarding(profile_manager: StudentProfileManager, gemini_client, user_id: str):
    st.title("🧑‍🎓 Construisons ton profil d'orientation")
    st.warning(
        "⚠️ **C'est important** : répondre à ces questions et charger tes 3 bulletins de "
        "terminale + ton relevé de notes du bac permet au chatbot d'analyser tes résultats "
        "et de te proposer une orientation vraiment personnalisée (adaptée à ton niveau, "
        "tes points forts et tes aspirations). Sans ça, les réponses resteront génériques.\n\n"
        "Tu peux toutefois passer cette étape si tu le souhaites et y revenir plus tard "
        "depuis le menu (« Modifier mon profil d'orientation »)."
    )

    col_titre, col_passer = st.columns([4, 1])
    with col_passer:
        if st.button("⏭️ Passer pour l'instant"):
            st.session_state.profil_onboarding_termine = True
            st.session_state.profil_a_ete_ignore = True
            st.rerun()

    profil = profile_manager.get_profile(user_id)

    etape = st.radio(
        "Étape", ["1️⃣ Aspirations et situation", "2️⃣ Bulletins et relevé de bac", "3️⃣ Analyse du profil"],
        horizontal=True,
    )

    # ---- Étape 1 : aspirations ----
    if etape.startswith("1"):
        with st.form("form_aspirations"):
            reponses = {}
            aspirations_existantes = profil.get("aspirations", {})
            for q in QUESTIONS_ASPIRATIONS:
                valeur_existante = aspirations_existantes.get(q["cle"])
                if q["type"] == "multiselect":
                    reponses[q["cle"]] = st.multiselect(q["label"], q["options"], default=valeur_existante or [])
                elif q["type"] == "select":
                    options = q["options"]
                    index = options.index(valeur_existante) if valeur_existante in options else 0
                    reponses[q["cle"]] = st.selectbox(q["label"], options, index=index)
                else:
                    reponses[q["cle"]] = st.text_input(q["label"], value=valeur_existante or "")
            if st.form_submit_button("💾 Enregistrer mes réponses"):
                profile_manager.save_aspirations(user_id, reponses)
                st.success("Réponses enregistrées ! Passe à l'étape 2 pour charger tes documents.")
                st.rerun()

    # ---- Étape 2 : upload des bulletins / relevé ----
    elif etape.startswith("2"):
        st.write("Charge chaque document au format PDF. Ils sont conservés de façon sécurisée dans la base de données.")
        documents_existants = (profile_manager.get_profile(user_id)).get("documents", {})
        for doc in DOCUMENTS_REQUIS:
            deja_charge = documents_existants.get(doc["cle"])
            statut = f"✅ déjà chargé : {deja_charge['nom_fichier']}" if deja_charge else "⏳ non chargé"
            st.write(f"**{doc['label']}** — {statut}")
            fichier = st.file_uploader(f"Charger : {doc['label']}", type=["pdf"], key=f"upload_{doc['cle']}")
            if fichier is not None:
                contenu_bytes = fichier.read()
                with st.spinner("Extraction du texte du document..."):
                    with open_temp_and_extract(contenu_bytes) as pages:
                        texte_extrait = "\n".join(t for _, t in pages)
                profile_manager.sauvegarder_document(user_id, doc["cle"], fichier.name, contenu_bytes, texte_extrait)
                st.success(f"{doc['label']} chargé avec succès.")
                st.rerun()

    # ---- Étape 3 : analyse académique ----
    else:
        manquants = profile_manager.documents_manquants(user_id)
        if manquants:
            st.warning("Charge d'abord tous les documents demandés à l'étape 2 avant de lancer l'analyse.")
        else:
            profil_actuel = profile_manager.get_profile(user_id)
            analyse_existante = profil_actuel.get("analyse_academique")
            if analyse_existante and not analyse_existante.get("erreur"):
                st.success("Ton profil académique a déjà été analysé :")
                st.json(analyse_existante)
                if st.button("🔄 Relancer l'analyse"):
                    with st.spinner("Analyse des bulletins en cours..."):
                        analyse = profile_manager.analyser_bulletins_avec_llm(gemini_client, user_id)
                    st.json(analyse)
            else:
                if st.button("🚀 Analyser mes bulletins et mon relevé de bac"):
                    with st.spinner("Analyse des bulletins en cours..."):
                        analyse = profile_manager.analyser_bulletins_avec_llm(gemini_client, user_id)
                    if analyse.get("erreur"):
                        st.error(analyse["erreur"])
                    else:
                        st.success("Analyse terminée !")
                        st.json(analyse)

    if profile_manager.profil_complet(user_id):
        st.markdown("---")
        st.success("✅ Ton profil est complet. Tu peux maintenant accéder au chat d'orientation.")
        if st.button("➡️ Accéder au chat d'orientation"):
            st.session_state.profil_onboarding_termine = True
            st.rerun()


class open_temp_and_extract:
    """Context manager utilitaire : écrit temporairement les octets d'un PDF
    uploadé sur disque pour réutiliser _extraire_texte_pdf, puis nettoie."""
    def __init__(self, contenu_bytes: bytes):
        self.contenu_bytes = contenu_bytes
        self.chemin_temp = None

    def __enter__(self):
        import tempfile
        fd, self.chemin_temp = tempfile.mkstemp(suffix=".pdf")
        with os.fdopen(fd, "wb") as f:
            f.write(self.contenu_bytes)
        return _extraire_texte_pdf(self.chemin_temp)

    def __exit__(self, exc_type, exc_val, exc_tb):
        if self.chemin_temp and os.path.exists(self.chemin_temp):
            try:
                os.remove(self.chemin_temp)
            except OSError:
                pass


# =========================================================================
# 🎨 Interface principale
# =========================================================================
def main():
    st.session_state.setdefault("authenticated", False)
    st.session_state.setdefault("user", None)
    st.session_state.setdefault("profil_onboarding_termine", False)
    st.session_state.setdefault("profil_a_ete_ignore", False)

    db = init_mongodb()
    if db is None:
        st.error("Impossible de se connecter à la base de données MongoDB.")
        return

    auth_manager = AuthManager(db)

    if not st.session_state.authenticated:
        show_login_page(auth_manager)
        return

    memory = ChatbotMemory(db)
    profile_manager = StudentProfileManager(db)
    gemini_client = init_gemini_client()
    user_id = st.session_state.user["user_id"]

    # --- Étape obligatoire : profil enrichi avant le chat ---
    profil_deja_complet = profile_manager.profil_complet(user_id)
    if not profil_deja_complet and not st.session_state.profil_onboarding_termine:
        with st.sidebar:
            st.write(f"👋 **{st.session_state.user['username']}**")
            if st.button("🚪 Se déconnecter"):
                st.session_state.authenticated = False
                st.session_state.user = None
                st.rerun()
        show_profile_onboarding(profile_manager, gemini_client, user_id)
        return

    st.title("🎓 Chatbot d'Orientation Universitaire")
    st.subheader(f"👋 Bienvenue {st.session_state.user['username']} !")

    if not profil_deja_complet:
        st.warning(
            "⚠️ Ton profil d'orientation n'est pas encore complet (aspirations et/ou "
            "bulletins manquants). Les recommandations resteront génériques tant que "
            "tu ne l'auras pas rempli. Tu peux le compléter à tout moment via "
            "« ✏️ Modifier mon profil d'orientation » dans le menu à gauche — "
            "c'est important pour être bien orienté(e)."
        )

    signature = _signature_dossier(DOCS_FOLDER)
    tables = charger_toutes_les_tables(DOCS_FOLDER, signature)
    if not tables:
        st.error(f"Aucune donnée structurée trouvée dans « {DOCS_FOLDER} ». "
                 f"Vérifie le chemin (variable d'environnement EDUCAT_FOLDER).")
        return
    chunks_docs = construire_corpus_documents(DOCS_FOLDER, signature)
    valeurs_connues = recuperer_toutes_les_valeurs_connues(tables)

    st.session_state.setdefault("current_session_id", memory.create_new_chat_session(user_id))
    st.session_state.setdefault("current_chat_history", [])

    with st.sidebar:
        st.header("🗂️ Menu")
        st.subheader("👤 Profil")
        st.write(f"**Nom :** {st.session_state.user['username']}")
        st.write(f"**Rôle :** {st.session_state.user['role']}")
        if st.button("✏️ Modifier mon profil d'orientation"):
            st.session_state.profil_onboarding_termine = False
            st.rerun()
        if st.button("🚪 Se déconnecter"):
            st.session_state.authenticated = False
            st.session_state.user = None
            st.session_state.profil_onboarding_termine = False
            st.rerun()

        st.markdown("---")
        st.subheader("💬 Sessions")
        if st.button("🆕 Nouveau Chat"):
            st.session_state.current_session_id = memory.create_new_chat_session(user_id)
            st.session_state.current_chat_history = []
            st.rerun()

        for session in memory.get_user_chat_sessions(user_id, 10):
            sid = str(session["_id"])
            titre = session.get("title", "Chat")
            if st.button(f"📄 {titre}", key=f"session_{sid}"):
                st.session_state.current_session_id = sid
                st.session_state.current_chat_history = []
                for conv in memory.get_session_conversations(sid):
                    st.session_state.current_chat_history.append({"role": "user", "content": conv["question"]})
                    st.session_state.current_chat_history.append({"role": "assistant", "content": conv["response"]})
                st.rerun()

        st.markdown("---")
        st.subheader("🔍 Consultés récemment")
        contexte = memory.get_user_context(user_id)
        for cle, etiquette in [("recent_universites", "Universités"), ("recent_filieres", "Filières"), ("recent_cites", "Cités")]:
            items = contexte.get(cle, [])
            if items:
                st.write(f"**{etiquette} :**")
                for i in items[:3]:
                    st.write(f"• {i['name']}")

        st.markdown("---")
        st.caption(f"📊 {len(tables)} table(s) de données chargée(s) — "
                    f"📄 {len(chunks_docs)} extrait(s) de documents indexé(s)")
        st.caption(f"📂 Source : {DOCS_FOLDER}")

    for message in st.session_state.current_chat_history:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    user_input = st.chat_input("Pose ta question sur les universités, filières, admissions ou cités...")

    if user_input:
        with st.chat_message("user"):
            st.write(user_input)

        with st.chat_message("assistant"):
            with st.spinner("Analyse croisée des sources en cours..."):
                reponse = traiter_question(
                    user_input, tables, chunks_docs, valeurs_connues, gemini_client,
                    memory, profile_manager, user_id, st.session_state.current_session_id,
                )
                st.write(reponse)

        st.session_state.current_chat_history.append({"role": "user", "content": user_input})
        st.session_state.current_chat_history.append({"role": "assistant", "content": reponse})

        if len(st.session_state.current_chat_history) == 2:
            titre = user_input[:50] + ("..." if len(user_input) > 50 else "")
            memory.update_session_title(st.session_state.current_session_id, titre)

    with st.expander("💡 Suggestions de questions"):
        st.write("• Quelles filières propose l'université Joseph Ki-Zerbo en sciences de la santé ?")
        st.write("• Quelles sont les conditions d'admission pour telle filière ?")
        st.write("• Quelles sont les débouchés réels d'une licence en économie et gestion au Burkina Faso ?")
        st.write("• Quelles cités universitaires existent à Koudougou ?")
        st.write("• Compare les universités publiques et privées pour une filière en informatique.")
        st.write("• Y a-t-il des bourses disponibles pour les filières scientifiques ?")


if __name__ == "__main__":
    main()
